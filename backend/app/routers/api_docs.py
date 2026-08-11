"""
世界 API 接口文档（程序员查看/下载用）——独立前缀，避免与 /worlds/{world_id} 冲突

与 view_api_doc 工具同源（data/world_api_docs/sections/*.md），AI 看到的文档这里都能看。

导出：md 直接下载；docx 走 pandoc（可选能力——容器装了 pandoc 才可用）。
"""
import io
import logging
import os
import re
import shutil
import subprocess
import zipfile

from fastapi import APIRouter, Depends, HTTPException, Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/kb", tags=["接口文档"])

DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ── docx 表格边框 + 斑马纹模板（pandoc 默认 reference.docx 的 Table 样式无边框无条纹，丑）──
_REF_DOCX_PATH = "/tmp/docx-ref-bordered-v3.docx"
_TBL_BORDERS = (
    "<w:tblBorders>"
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    "</w:tblBorders>"
)
# 行带斑马纹（Word 条件格式 band1Horz，浅灰底）
_TBL_BANDING = (
    '<w:tblStylePr w:type="band1Horz"><w:tcPr>'
    '<w:shd w:val="clear" w:color="auto" w:fill="F2F2F2"/>'
    "</w:tcPr></w:tblStylePr>"
)


# ── docx 产品色美化（violet 品牌色：主 #7C3AED / 浅紫 #EDE9FE / 紫调灰 #F8F7FC / 代码底 #F6F5FA）──
_BRAND_VIOLET = "7C3AED"
_HEADER_FILL = "EDE9FE"
_ZEBRA_FILL = "F8F7FC"
_CODE_FILL = "F6F5FA"


def _inject_style_rpr(xml: str, style_id: str, rpr_xml: str) -> str:
    """给指定 styleId 的样式块注入 rPr 内容（幂等）"""
    marker = f"<w:styleId>{style_id}</w:styleId>"
    pos = xml.find(marker)
    if pos < 0:
        return xml
    start = xml.rfind("<w:style ", 0, pos)
    end = xml.find("</w:style>", pos)
    if start < 0 or end < 0:
        return xml
    block = xml[start:end]
    if rpr_xml in block:
        return xml
    if "<w:rPr>" in block:
        new_block = block.replace("</w:rPr>", rpr_xml + "</w:rPr>", 1)
    else:
        new_block = block.replace("</w:style>", "<w:rPr>" + rpr_xml + "</w:rPr></w:style>", 1)
    return xml[:start] + new_block + xml[end:]


def _inject_style_ppr(xml: str, style_id: str, ppr_xml: str) -> str:
    """给指定 styleId 的样式块注入 pPr 内容（仅当块内已有 pPr；幂等）"""
    marker = f"<w:styleId>{style_id}</w:styleId>"
    pos = xml.find(marker)
    if pos < 0:
        return xml
    start = xml.rfind("<w:style ", 0, pos)
    end = xml.find("</w:style>", pos)
    if start < 0 or end < 0:
        return xml
    block = xml[start:end]
    if ppr_xml in block or "<w:pPr>" not in block:
        return xml
    new_block = block.replace("</w:pPr>", ppr_xml + "</w:pPr>", 1)
    return xml[:start] + new_block + xml[end:]


def _ensure_ref_docx() -> str:
    """生成带表格边框的 pandoc reference docx（基于默认模板改 styles.xml，缓存 /tmp）"""
    if os.path.exists(_REF_DOCX_PATH):
        return _REF_DOCX_PATH
    try:
        proc = subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            capture_output=True, timeout=15,
        )
        if proc.returncode != 0:
            return ""
        zin = zipfile.ZipFile(io.BytesIO(proc.stdout))
        styles = zin.read("word/styles.xml").decode("utf-8")

        def add_borders(m: re.Match) -> str:
            block = m.group(0)
            # 已注入过（边框 + 斑马纹都在）则跳过；pandoc 默认模板自带 firstRow 的 tblStylePr，不能以此判断
            if "<w:tblBorders>" in block and 'w:type="band1Horz"' in block:
                return block
            if "<w:tblPr" in block and "</w:tblPr>" in block:
                # 边框进 tblPr，斑马纹紧跟 tblPr 之后（Word schema 顺序：tblPr → tblStylePr）
                return block.replace(
                    "</w:tblPr>",
                    _TBL_BORDERS + "</w:tblPr>" + _TBL_BANDING,
                    1,
                )
            return block

        # 默认表格样式：w:type="table" 且 w:default="1"（pandoc reference.docx 的 Table）
        styles2 = re.sub(
            r'<w:style [^>]*w:type="table"[^>]*w:default="1".*?</w:style>',
            add_borders, styles, flags=re.S,
        )
        if styles2 == styles:
            # 兜底：按 styleId="Table" 匹配
            styles2 = re.sub(
                r'<w:style [^>]*w:type="table"[^>]*w:styleId="Table".*?</w:style>',
                add_borders, styles, flags=re.S,
            )
        # ── 产品色美化：标题/链接品牌紫，代码块 Consolas 9pt + 紫调底纹 ──
        for sid in ("Heading1", "Heading2", "Heading3"):
            styles2 = _inject_style_rpr(styles2, sid, f'<w:color w:val="{_BRAND_VIOLET}"/>')
        styles2 = _inject_style_rpr(styles2, "Hyperlink", f'<w:color w:val="{_BRAND_VIOLET}"/>')
        styles2 = _inject_style_rpr(
            styles2, "SourceCode",
            '<w:rFonts w:ascii="Consolas" w:hAnsi="Consolas" w:eastAsia="Consolas"/>'
            '<w:sz w:val="18"/><w:color w:val="334155"/>',
        )
        styles2 = _inject_style_ppr(styles2, "SourceCode", f'<w:shd w:val="clear" w:color="auto" w:fill="{_CODE_FILL}"/>')
        styles2 = _inject_style_rpr(
            styles2, "VerbatimChar",
            '<w:rFonts w:ascii="Consolas" w:hAnsi="Consolas" w:eastAsia="Consolas"/>',
        )
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                content = zin.read(item.filename)
                if item.filename == "word/styles.xml":
                    content = styles2.encode("utf-8")
                zout.writestr(item, content)
        with open(_REF_DOCX_PATH, "wb") as f:
            f.write(buf.getvalue())
        return _REF_DOCX_PATH
    except Exception:
        logger.warning("生成 docx 边框模板失败，退回无模板", exc_info=True)
        return ""

# pandoc 在线安装状态（内存态；重启后 pandoc 已装则 docx_available 直接为 True）
_INSTALL_STATE: dict = {"running": False, "ok": False, "error": "", "started_at": ""}


def _post_process_docx(data: bytes) -> bytes:
    """docx 后处理：表头浅紫底+深紫加粗字；偶数行紫调斑马纹（硬编码单元格 shd，任何查看器都显示）"""
    try:
        from io import BytesIO

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor

        doc = Document(BytesIO(data))
        for tbl in doc.tables:
            for idx, row in enumerate(tbl.rows):
                if idx == 0:
                    fill = _HEADER_FILL
                elif idx % 2 == 1:
                    fill = _ZEBRA_FILL
                else:
                    continue
                seen = set()
                for cell in row.cells:
                    if id(cell._tc) in seen:
                        continue  # 横向合并单元格去重
                    seen.add(id(cell._tc))
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), fill)
                    tc_pr.append(shd)
                    if idx == 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
                                run.font.bold = True
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        logger.warning("docx 后处理失败，返回原 docx", exc_info=True)
        return data


def docx_available() -> bool:
    """pandoc 是否可用（可选依赖：容器装了/在线装完才有 docx 导出）"""
    return shutil.which("pandoc") is not None


async def _install_pandoc_background() -> None:
    """后台安装 pandoc（apt-get；管理员触发，进度靠轮询 status）"""
    from datetime import datetime, timezone
    _INSTALL_STATE["running"] = True
    _INSTALL_STATE["error"] = ""
    _INSTALL_STATE["started_at"] = datetime.now(timezone.utc).isoformat()
    try:
        import asyncio
        proc = await asyncio.create_subprocess_exec(
            "apt-get", "update", "-qq",
            stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL,
        )
        await proc.wait()
        proc = await asyncio.create_subprocess_exec(
            "apt-get", "install", "-y", "-qq", "--no-install-recommends", "pandoc",
            stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL,
        )
        await proc.wait()
        _INSTALL_STATE["ok"] = proc.returncode == 0
        _INSTALL_STATE["error"] = "" if proc.returncode == 0 else f"apt 安装失败（code {proc.returncode}）"
    except Exception as e:
        _INSTALL_STATE["ok"] = False
        _INSTALL_STATE["error"] = str(e)
    finally:
        _INSTALL_STATE["running"] = False


@router.get("")
async def list_api_docs(current_user: dict = Depends(get_current_user)):
    """分区列表（id / 标题 / 区介绍）"""
    from app.services.world.world_api_docs import SECTIONS
    return {"sections": SECTIONS}


@router.get("/status")
async def export_status(
    current_user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """导出能力状态：docx 是否可用 + 当前用户是否管理员（未装时管理员看安装提示，普通用户不显示）"""
    from app.models.user import User
    u = await db.get(User, current_user["user_id"])
    is_admin = bool(u and u.role == "admin")
    return {
        "docx_available": docx_available(),
        "is_admin": is_admin,
        "installing": _INSTALL_STATE["running"],
        "install_error": _INSTALL_STATE["error"] or None,
    }


@router.post("/install")
async def install_pandoc(
    current_user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """管理员在线安装 pandoc（apt-get，后台执行；轮询 export/status 看 installing/ok）"""
    from app.models.user import User
    u = await db.get(User, current_user["user_id"])
    if not (u and u.role == "admin"):
        raise HTTPException(status_code=403, detail="仅管理员可安装")
    if docx_available():
        return {"docx_available": True, "installing": False}
    if _INSTALL_STATE["running"]:
        return {"docx_available": False, "installing": True}
    import asyncio
    asyncio.create_task(_install_pandoc_background())
    return {"docx_available": False, "installing": True}


class DocxExportRequest(BaseModel):
    """md 内容 → docx（前端组织好内容（单分区/全部合并），后端只管转换）"""
    md: str
    filename: str = "api-doc.docx"


@router.post("/convert")
async def export_docx(
    req: DocxExportRequest,
    current_user: dict = Depends(get_current_user),
):
    """pandoc 把 md 转 docx 下载（未装 pandoc 返回 400，前端据此隐藏 docx 选项）"""
    if not docx_available():
        raise HTTPException(status_code=400, detail="docx 导出未安装（需要 pandoc）")
    try:
        cmd = [
            "pandoc", "-f", "markdown", "-t", "docx",
            "--toc", "--toc-depth=2",
            "--highlight-style=tango",  # 代码块语法高亮（tango 配色，Word 生态最顺眼）
        ]
        ref = _ensure_ref_docx()
        if ref:
            cmd.append(f"--reference-doc={ref}")
        result = subprocess.run(
            cmd,
            input=req.md.encode("utf-8"), capture_output=True, timeout=30,
        )
    except Exception as e:
        logger.error(f"pandoc 转换失败: {e}")
        raise HTTPException(status_code=500, detail=f"docx 转换失败: {e}")
    if result.returncode != 0:
        raise HTTPException(status_code=500, detail=f"pandoc 错误: {result.stderr.decode(errors='replace')[:200]}")
    filename = req.filename if req.filename.endswith(".docx") else req.filename + ".docx"
    return Response(
        _post_process_docx(result.stdout),
        media_type=DOCX_MEDIA,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{section_id}")
async def get_api_doc(section_id: str, current_user: dict = Depends(get_current_user)):
    """读取某个分区的接口文档内容（防路径穿越：仅注册表内 id）"""
    from app.services.world.world_api_docs import view_section
    try:
        return view_section(section_id)
    except (ValueError, FileNotFoundError) as e:
        raise HTTPException(status_code=404, detail=str(e))
